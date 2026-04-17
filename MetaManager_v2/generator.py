import os
import copy
import re
import shutil
import subprocess
from datetime import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import lxml.etree as etree


class KPGenerator:
    def __init__(self, templates_dir="templates"):
        self.templates_dir = os.path.join(os.path.dirname(os.path.abspath(__file__)), templates_dir)

    def get_template_name(self, branch, volume, smr_type):
        branch_map = {"хоз.пит": "hozpit", "техническая лицензия": "tech"}
        volume_map = {
            "до 100": "100",
            "100-500": "100_500",
            "500+": "500plus",
            "500+ с переоценкой запасов": "500pluspere"
        }
        smr_map = {"с смр": "with_smr", "без смр": "no_smr"}
        template_filename = f"{branch_map[branch]}_{volume_map[volume]}_{smr_map[smr_type]}.docx"
        return os.path.join(self.templates_dir, template_filename)

    def safe_float(self, val):
        if isinstance(val, (int, float)):
            return float(val)
        try:
            clean_val = str(val).replace(' ', '').replace(',', '.')
            return float(clean_val)
        except (ValueError, TypeError):
            return 0.0

    def format_price(self, val):
        return f"{int(round(val)):,}".replace(',', ' ')

    # -----------------------------------------------------------------------
    # Утилиты для работы с runs в параграфах
    # -----------------------------------------------------------------------

    def _para_get_text(self, para):
        """Возвращает полный текст параграфа из всех runs."""
        return ''.join(r.text for r in para.runs)

    def _para_set_text(self, para, new_text):
        """
        Устанавливает текст параграфа, сохраняя форматирование первого run.
        Все лишние runs удаляются.
        """
        runs = para.runs
        if runs:
            # Сохраняем форматирование первого run
            runs[0].text = new_text
            # Удаляем остальные runs
            for run in runs[1:]:
                run._element.getparent().remove(run._element)
        else:
            para.add_run(new_text)

    def _para_replace(self, para, replacements):
        """
        Заменяет плейсхолдеры в параграфе.
        replacements: dict {old: new}
        """
        text = self._para_get_text(para)
        changed = False
        for old, new in replacements.items():
            if old in text:
                text = text.replace(old, new)
                changed = True
        if changed:
            self._para_set_text(para, text)

    def _cell_get_text(self, cell):
        """Возвращает полный текст ячейки."""
        return cell.text

    def _cell_replace(self, cell, replacements):
        """Заменяет плейсхолдеры во всех параграфах ячейки."""
        for para in cell.paragraphs:
            self._para_replace(para, replacements)

    def _row_replace(self, row, replacements):
        """Заменяет плейсхолдеры во всех ячейках строки."""
        for cell in row.cells:
            self._cell_replace(cell, replacements)

    def _row_get_text(self, row):
        """Возвращает полный текст строки."""
        return ' '.join(cell.text for cell in row.cells)

    def _cell_set_text(self, cell, new_text):
        """Устанавливает текст в первом параграфе ячейки."""
        para = cell.paragraphs[0] if cell.paragraphs else None
        if para is None:
            return
        self._para_set_text(para, new_text)
        # Удаляем лишние параграфы
        for p in cell.paragraphs[1:]:
            p._element.getparent().remove(p._element)

    def _ensure_cell_border(self, cell, side, size='8'):
        """Гарантирует наличие конкретной границы у ячейки, не ломая остальные стили."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            tc_borders = OxmlElement('w:tcBorders')
            tc_pr.append(tc_borders)

        border = tc_borders.find(qn(f'w:{side}'))
        if border is None:
            border = OxmlElement(f'w:{side}')
            tc_borders.append(border)

        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), size)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')

    def _remove_cell_border(self, cell, side):
        """Убирает конкретную границу ячейки, чтобы не получить двойную линию."""
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        tc_borders = tc_pr.find(qn('w:tcBorders'))
        if tc_borders is None:
            return
        border = tc_borders.find(qn(f'w:{side}'))
        if border is not None:
            tc_borders.remove(border)

    def _set_exact_cell_borders(self, cell, sides, size='8'):
        """
        Полностью задаёт границы ячейки по списку sides, удаляя конфликтующие настройки.
        sides: iterable из {'top','left','bottom','right'}.
        """
        tc = cell._tc
        tc_pr = tc.get_or_add_tcPr()
        old_borders = tc_pr.find(qn('w:tcBorders'))
        if old_borders is not None:
            tc_pr.remove(old_borders)

        tc_borders = OxmlElement('w:tcBorders')
        for side in sides:
            border = OxmlElement(f'w:{side}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), size)
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), '000000')
            tc_borders.append(border)
        tc_pr.append(tc_borders)

    def _set_cell_paragraph_alignment(self, cell, alignment):
        """Устанавливает выравнивание для всех параграфов ячейки."""
        for para in cell.paragraphs:
            para.alignment = alignment

    def normalize_table_borders(self, doc):
        """
        Точечно исправляет проблемные границы таблицы, сохраняя шаблонный стиль.
        """
        if not doc.tables:
            return
        table = doc.tables[0]

        # Колонка "Срок" обычно индекс 2, "Результат" - индекс 3.
        # Добавляем границы между ними, если таблица имеет нужные столбцы.
        for row in table.rows:
            row_text = self._row_get_text(row).lower()
            if len(row.cells) >= 4 and 'итого без ндс' not in row_text and 'итого с ндс' not in row_text:
                # Ставим границу только с одной стороны стыка (иначе в PDF выглядит толще).
                self._ensure_cell_border(row.cells[2], 'right', size='6')
                self._remove_cell_border(row.cells[3], 'left')

        # Для строк итогов гарантируем верх/низ, чтобы убрать "пустоты".
        total_rows = []
        for row_idx, row in enumerate(table.rows):
            row_text = self._row_get_text(row).lower()
            cell0_text = row.cells[0].text.strip() if row.cells else ''
            cell1_text = row.cells[1].text.strip().lower() if len(row.cells) > 1 else ''

            # Для цветных строк "Этап N..." гарантируем верхнюю границу.
            if cell0_text.isdigit() and 'этап' in cell1_text:
                for cell in row.cells:
                    self._ensure_cell_border(cell, 'top', size='8')
                # Дублируем как нижнюю границу предыдущей строки для стабильного рендера в PDF.
                if row_idx > 0:
                    prev_row = table.rows[row_idx - 1]
                    for prev_cell in prev_row.cells:
                        self._ensure_cell_border(prev_cell, 'bottom', size='8')

            if 'итого без ндс' in row_text or 'итого с ндс' in row_text:
                total_rows.append(row_idx)

        # Точное формирование блока итогов без двойных/пропадающих линий в PDF.
        for row_idx in total_rows:
            row = table.rows[row_idx]
            if len(row.cells) < 2:
                continue
            label_cell = row.cells[-2]
            value_cell = row.cells[-1]

            # Одна граница слева у блока: ставим справа у предыдущей ячейки.
            if len(row.cells) >= 3:
                prev_cell = row.cells[-3]
                self._ensure_cell_border(prev_cell, 'right', size='8')
                self._remove_cell_border(label_cell, 'left')

            # Явный разделитель "подпись / сумма" + верх/низ обеих итоговых ячеек.
            self._set_exact_cell_borders(label_cell, ('top', 'bottom', 'right'), size='8')
            self._set_exact_cell_borders(value_cell, ('top', 'bottom', 'right'), size='8')

    # -----------------------------------------------------------------------
    # Операции со строками таблицы
    # -----------------------------------------------------------------------

    def _copy_row_after(self, table, source_row_idx):
        """
        Копирует строку таблицы и вставляет её сразу после source_row_idx.
        Возвращает новую строку.
        """
        source_row = table.rows[source_row_idx]
        new_tr = copy.deepcopy(source_row._tr)
        source_row._tr.addnext(new_tr)
        return table.rows[source_row_idx + 1]

    def _delete_row(self, table, row_idx):
        """Удаляет строку таблицы по индексу."""
        row = table.rows[row_idx]
        row._tr.getparent().remove(row._tr)

    # -----------------------------------------------------------------------
    # Обработка шаблонов с СМР
    # -----------------------------------------------------------------------

    def process_smr_template(self, doc, data):
        """
        Обрабатывает шаблон с СМР:
        - Заменяет jinja2-теги на реальные данные
        - Правильно дублирует строки скважин и насосов под оригинальными
        - Удаляет строку БМЗ если не нужна
        - Обновляет нумерацию строк этапа 3
        """
        table = doc.tables[0]
        wells_count = int(data.get('wells_count', 1) or 1)
        wells_price = self.safe_float(data.get('wells_price', 0))
        pump_price = self.safe_float(data.get('pump_price', 0))
        bmz_price = self.safe_float(data.get('bmz_price', 0))
        include_wells = data.get('include_wells', True)
        include_pump = data.get('include_pump', True)
        include_bmz = data.get('include_bmz', True)
        wells_design = data.get('wells_design', '')
        wells_depth = data.get('wells_depth', '')
        bmz_size = data.get('bmz_size', '')

        # Находим строки с jinja2-тегами
        wells_row_idx = -1
        pumps_row_idx = -1
        bmz_row_idx = -1

        for i, row in enumerate(table.rows):
            row_text = self._row_get_text(row)
            if '{% for w in wells' in row_text:
                wells_row_idx = i
            elif '{% for p in pumps' in row_text:
                pumps_row_idx = i
            elif '{% if include_bmz' in row_text:
                bmz_row_idx = i

        # --- Обрабатываем строку скважин ---
        if wells_row_idx >= 0:
            wells_row = table.rows[wells_row_idx]

            if include_wells:
                # Сначала создаём копии строки для скважин 2, 3, ...
                # (копируем ДО замены, чтобы копии содержали плейсхолдеры)
                for i in range(2, wells_count + 1):
                    # Вставляем после предыдущей скважины
                    insert_after = wells_row_idx + (i - 2)
                    self._copy_row_after(table, insert_after)
                    # Обновляем индексы
                    pumps_row_idx += 1
                    bmz_row_idx += 1

                # Теперь заполняем все строки скважин
                for i in range(1, wells_count + 1):
                    row_idx = wells_row_idx + (i - 1)
                    row = table.rows[row_idx]
                    replacements = {
                        '{% for w in wells %}': '',
                        '{% endfor %}': '',
                        '{% for w in wells%}': '',
                        '{%endfor%}': '',
                        '{{ w.design }}': wells_design,
                        '{{w.design}}': wells_design,
                        '{{ w.depth }}': str(wells_depth),
                        '{{w.depth}}': str(wells_depth),
                        '{{ w.price }}': self.format_price(wells_price),
                        '{{w.price}}': self.format_price(wells_price),
                    }
                    self._row_replace(row, replacements)
                    # Устанавливаем номер
                    self._cell_set_text(row.cells[0], f'3.{i}')
            else:
                # Удаляем строку скважин
                self._delete_row(table, wells_row_idx)
                pumps_row_idx -= 1
                bmz_row_idx -= 1

        # --- Обрабатываем строку насосов ---
        if pumps_row_idx >= 0:
            pumps_row = table.rows[pumps_row_idx]

            if include_pump:
                # Создаём копии для насосов 2, 3, ...
                for i in range(2, wells_count + 1):
                    insert_after = pumps_row_idx + (i - 2)
                    self._copy_row_after(table, insert_after)
                    bmz_row_idx += 1

                # Заполняем все строки насосов
                pump_base_num = (wells_count if include_wells else 0) + 1
                for i in range(1, wells_count + 1):
                    row_idx = pumps_row_idx + (i - 1)
                    row = table.rows[row_idx]
                    replacements = {
                        '{% for p in pumps %}': '',
                        '{% endfor %}': '',
                        '{% for p in pumps%}': '',
                        '{%endfor%}': '',
                        '{{ p.price }}': self.format_price(pump_price),
                        '{{p.price}}': self.format_price(pump_price),
                    }
                    self._row_replace(row, replacements)
                    pump_num = pump_base_num + (i - 1)
                    self._cell_set_text(row.cells[0], f'3.{pump_num}')
            else:
                # Удаляем строку насосов
                self._delete_row(table, pumps_row_idx)
                bmz_row_idx -= 1

        # --- Обрабатываем строку БМЗ ---
        if bmz_row_idx >= 0:
            bmz_row = table.rows[bmz_row_idx]

            if include_bmz:
                replacements = {
                    '{% if include_bmz %}': '',
                    '{% endif %}': '',
                    '{% if include_bmz%}': '',
                    '{%endif%}': '',
                    '{{ bmz_size }}': bmz_size,
                    '{{bmz_size}}': bmz_size,
                    '{{ bmz_price }}': self.format_price(bmz_price),
                    '{{bmz_price}}': self.format_price(bmz_price),
                }
                self._row_replace(bmz_row, replacements)
                # Обновляем номер БМЗ
                bmz_num = self._calc_bmz_num(include_wells, include_pump, wells_count)
                self._cell_set_text(bmz_row.cells[0], f'3.{bmz_num}')
            else:
                # Удаляем строку БМЗ
                self._delete_row(table, bmz_row_idx)
                bmz_row_idx = -1

        # --- Перенумеровываем статические строки этапа 3 ---
        self._renumber_stage3_rows(table, include_wells, include_pump, include_bmz, wells_count)

    def _calc_bmz_num(self, include_wells, include_pump, wells_count):
        """Вычисляет номер строки БМЗ в этапе 3."""
        n = 0
        if include_wells:
            n += wells_count
        if include_pump:
            n += wells_count
        return n + 1

    def _renumber_stage3_rows(self, table, include_wells, include_pump, include_bmz, wells_count):
        """
        Перенумеровывает статические строки этапа 3 после динамических.
        """
        # Считаем сколько строк занимают скважины + насосы + БМЗ
        dynamic_count = 0
        if include_wells:
            dynamic_count += wells_count
        if include_pump:
            dynamic_count += wells_count
        if include_bmz:
            dynamic_count += 1

        # Находим строку "3 | Этап 3..."
        stage3_header_idx = -1
        for i, row in enumerate(table.rows):
            cell0 = row.cells[0].text.strip()
            cell1 = row.cells[1].text.strip()
            if cell0 == '3' and 'Этап 3' in cell1:
                stage3_header_idx = i
                break

        if stage3_header_idx < 0:
            return

        # Находим строку "4 | Этап 4..."
        stage4_header_idx = -1
        for i, row in enumerate(table.rows):
            cell0 = row.cells[0].text.strip()
            cell1 = row.cells[1].text.strip()
            if cell0 == '4' and 'Этап 4' in cell1:
                stage4_header_idx = i
                break

        if stage4_header_idx < 0:
            return

        # Строки после динамических (начиная с stage3_header_idx + 1 + dynamic_count)
        static_start = stage3_header_idx + 1 + dynamic_count
        new_num = dynamic_count + 1

        for row_idx in range(static_start, stage4_header_idx):
            row = table.rows[row_idx]
            cell0_text = row.cells[0].text.strip()
            if cell0_text.startswith('3.'):
                self._cell_set_text(row.cells[0], f'3.{new_num}')
                new_num += 1

    # -----------------------------------------------------------------------
    # Обновление хим.анализа и ОФР
    # -----------------------------------------------------------------------

    def update_dynamic_fields(self, doc, wells_count):
        """
        Обновляет поля хим.анализа и ОФР в зависимости от количества скважин.
        """
        table = doc.tables[0]

        for i, row in enumerate(table.rows):
            row_text = self._row_get_text(row)

            # Строка с хим.анализом
            if 'Расширенный химический анализ' in row_text and '54 компонента' in row_text:
                # Получаем текущую цену из ячейки 5 (индекс 4)
                price_text = row.cells[4].text.strip().replace(' ', '').replace(',', '.')
                try:
                    price_per_well = float(price_text)
                except ValueError:
                    price_per_well = 120000

                # Обновляем "Результат работ" (ячейка 4, индекс 3)
                self._cell_set_text(row.cells[3], f'полный хим. анализ ({wells_count} шт)')
                # Обновляем сумму
                self._cell_set_text(row.cells[4], self.format_price(price_per_well * wells_count))

            # Строка с ОФР
            elif 'Опытно-фильтрационные работы' in row_text:
                # Определяем количество дней из названия строки
                days = 5  # дефолт
                if '5 дней' in row_text or '5 суток' in row_text:
                    days = 5
                elif '3 дня' in row_text or '3 суток' in row_text or '3 дней' in row_text:
                    days = 3

                # Получаем текущую цену
                price_text = row.cells[4].text.strip().replace(' ', '').replace(',', '.')
                try:
                    price_per_well = float(price_text)
                except ValueError:
                    price_per_well = 150000

                # Обновляем "Результат работ"
                self._cell_set_text(row.cells[3], f'Журнал откачки ({days}*{wells_count}) суток')
                # Обновляем сумму
                self._cell_set_text(row.cells[4], self.format_price(price_per_well * wells_count))

            # Строка с геофизическими исследованиями
            elif 'Геофизические исследования в скважине' in row_text:
                # Получаем текущую цену за 1 скважину
                price_text = row.cells[4].text.strip().replace(' ', '').replace(',', '.')
                try:
                    price_per_well = float(price_text)
                except ValueError:
                    price_per_well = 0

                # Обновляем "Результат работ" (если есть столбец результата)
                if len(row.cells) > 3:
                    self._cell_set_text(row.cells[3], f'Геофизическое заключение ({wells_count} шт)')

                # Обновляем сумму
                self._cell_set_text(row.cells[4], self.format_price(price_per_well * wells_count))

    def add_pir_stage(self, doc, pir_count, pir_price):
        """
        Добавляет новый этап ПИР перед текущим первым этапом и сдвигает нумерацию остальных.
        """
        if not doc.tables:
            return
        table = doc.tables[0]

        first_stage_idx = -1
        first_substage_idx = -1
        for i, row in enumerate(table.rows):
            cell0 = row.cells[0].text.strip() if len(row.cells) > 0 else ""
            cell1 = row.cells[1].text.strip() if len(row.cells) > 1 else ""
            if first_stage_idx < 0 and cell0 == "1" and "Этап" in cell1:
                first_stage_idx = i
                continue
            if first_stage_idx >= 0 and first_substage_idx < 0 and cell0.startswith("1."):
                first_substage_idx = i
                break

        if first_stage_idx < 0 or first_substage_idx < 0:
            return

        source_header_tr = table.rows[first_stage_idx]._tr
        source_item_tr = table.rows[first_substage_idx]._tr

        new_header_tr = copy.deepcopy(source_header_tr)
        source_header_tr.addprevious(new_header_tr)
        new_item_tr = copy.deepcopy(source_item_tr)
        new_header_tr.addnext(new_item_tr)

        header_idx = first_stage_idx
        item_idx = first_stage_idx + 1

        header_row = table.rows[header_idx]
        item_row = table.rows[item_idx]
        self._cell_set_text(header_row.cells[0], "1")
        self._cell_set_text(header_row.cells[1], "Этап 1. Проект ВЗУ (стадии П и Р)")
        if len(header_row.cells) > 2:
            self._cell_set_text(header_row.cells[2], "3 мес")
            self._set_cell_paragraph_alignment(header_row.cells[2], WD_PARAGRAPH_ALIGNMENT.CENTER)

        pir_total = self.safe_float(pir_count) * self.safe_float(pir_price)
        self._cell_set_text(item_row.cells[0], "1.1")
        self._cell_set_text(
            item_row.cells[1],
            "Разработка проекта ВЗУ ( Стадии П и Р: ПЗ,ПЗУ, АР, КР, ИОС1, ИОС2, ИОС3, "
            "ИОС4, ИОС5, ТХ, ПОС, ООС, ПБ, ТБЭО, ГП, АР, КР (КЖ, КМ), ЭО, НВ, НК, ОВ, СС, ТХ)"
        )
        if len(item_row.cells) > 2:
            self._cell_set_text(item_row.cells[2], "3 мес")
        if len(item_row.cells) > 3:
            self._cell_set_text(item_row.cells[3], "Проект ВЗУ")
        if len(item_row.cells) > 4:
            self._cell_set_text(item_row.cells[4], self.format_price(pir_total))

        self._shift_stage_numbers(table, start_row=item_idx + 1)

    def _shift_stage_numbers(self, table, start_row):
        """Сдвигает номера этапов/подэтапов на +1, начиная с указанной строки."""
        for row_idx in range(start_row, len(table.rows)):
            row = table.rows[row_idx]
            if not row.cells:
                continue

            cell0_text = row.cells[0].text.strip()
            match = re.match(r'^(\d+)(?:\.(\d+))?$', cell0_text)
            if not match:
                continue

            stage_num = int(match.group(1)) + 1
            sub_num = match.group(2)
            if sub_num is None:
                old_stage_num = int(match.group(1))
                self._cell_set_text(row.cells[0], str(stage_num))
                if len(row.cells) > 1:
                    stage_title = row.cells[1].text
                    stage_title = re.sub(
                        r'Этап\s+\d+',
                        f'Этап {stage_num}',
                        stage_title
                    )
                    self._cell_set_text(row.cells[1], stage_title)
                if len(row.cells) > 4:
                    stage_sum_text = row.cells[4].text
                    stage_sum_text = stage_sum_text.replace(
                        f"{{{{ stage{old_stage_num}_total }}}}",
                        f"{{{{ stage{stage_num}_total }}}}"
                    )
                    stage_sum_text = stage_sum_text.replace(
                        f"{{{{stage{old_stage_num}_total}}}}",
                        f"{{{{stage{stage_num}_total}}}}"
                    )
                    self._cell_set_text(row.cells[4], stage_sum_text)
            else:
                self._cell_set_text(row.cells[0], f'{stage_num}.{sub_num}')

    def increase_total_duration(self, doc, months_delta):
        """
        Увеличивает общий срок работ (первая найденная ячейка с шаблоном 'N мес').
        """
        if not doc.tables:
            return
        table = doc.tables[0]

        for row in table.rows:
            for cell in row.cells:
                text = cell.text
                match = re.search(r'(\d+)\s*мес', text, flags=re.IGNORECASE)
                if not match:
                    continue

                current_months = int(match.group(1))
                updated_months = current_months + int(months_delta)
                updated_text = re.sub(
                    r'(\d+)\s*мес',
                    f'{updated_months} мес',
                    text,
                    count=1,
                    flags=re.IGNORECASE
                )
                self._cell_set_text(cell, updated_text)
                return

    def _export_pdf_windows_word(self, docx_path, pdf_path):
        """
        Экспорт в PDF через Microsoft Word COM.
        Обычно даёт наиболее точное совпадение с отображением в Word.
        """
        import win32com.client  # type: ignore

        word = win32com.client.DispatchEx('Word.Application')
        word.Visible = False
        doc = None
        try:
            doc = word.Documents.Open(docx_path)
            # 17 = wdExportFormatPDF
            # OpenAfterExport=False, OptimizeFor=0 (print)
            doc.ExportAsFixedFormat(pdf_path, 17, False, 0)
        finally:
            if doc is not None:
                doc.Close(False)
            word.Quit()

    # -----------------------------------------------------------------------
    # Замена простых плейсхолдеров
    # -----------------------------------------------------------------------

    def replace_placeholders(self, doc, context):
        """Заменяет плейсхолдеры {{ var }} в параграфах и таблицах документа."""
        replacements = {'{{ ' + k + ' }}': str(v) for k, v in context.items()}
        # Добавляем варианты без пробелов
        replacements.update({'{{' + k + '}}': str(v) for k, v in context.items()})

        # В параграфах документа
        for para in doc.paragraphs:
            self._para_replace(para, replacements)

        # В таблицах
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    self._cell_replace(cell, replacements)

    # -----------------------------------------------------------------------
    # Вычисление итогов
    # -----------------------------------------------------------------------

    def calculate_totals(self, doc):
        """
        Вычисляет итоговые суммы по этапам, суммируя числовые значения в столбце сумм.
        """
        table = doc.tables[0]

        stage_totals = {}
        current_stage = None

        for i, row in enumerate(table.rows):
            cell0 = row.cells[0].text.strip()
            cell4 = row.cells[4].text.strip() if len(row.cells) > 4 else ''

            # Определяем заголовок этапа
            if cell0.isdigit() and len(row.cells) > 1:
                cell1 = row.cells[1].text.strip()
                if 'Этап' in cell1:
                    current_stage = int(cell0)
                    if current_stage not in stage_totals:
                        stage_totals[current_stage] = 0
                    continue

            # Строка подпункта
            if current_stage and '.' in cell0:
                # Пытаемся получить цену (убираем пробелы и плейсхолдеры)
                price_text = cell4.replace(' ', '').replace(',', '.')
                if '{{' not in price_text and price_text:
                    try:
                        price = float(price_text)
                        stage_totals[current_stage] = stage_totals.get(current_stage, 0) + price
                    except ValueError:
                        pass

        grand_total = sum(stage_totals.values())
        totals = {}
        for stage_num, total in stage_totals.items():
            totals[f'stage{stage_num}_total'] = total
        totals['subtotal'] = grand_total
        totals['vat'] = grand_total * 0.05
        totals['total_with_vat'] = grand_total + grand_total * 0.05

        return totals

    # -----------------------------------------------------------------------
    # Основной метод
    # -----------------------------------------------------------------------

    def create_kp(self, data):
        template_path = self.get_template_name(data['branch'], data['volume'], data['smr_type'])

        if not os.path.exists(template_path):
            raise FileNotFoundError(f"Шаблон не найден: {template_path}")

        doc = Document(template_path)

        wells_count = int(data.get('wells_count', 1) or 1)
        is_smr = data.get('smr_type') == 'с смр'

        # Шаг 1: Обработка динамических строк СМР
        if is_smr:
            self.process_smr_template(doc, data)

        # Шаг 2: Обновляем хим.анализ и ОФР (для всех шаблонов)
        self.update_dynamic_fields(doc, wells_count)

        # Шаг 2.1: Добавляем этап ПИР (если выбран)
        if data.get('include_pir'):
            self.add_pir_stage(
                doc,
                data.get('pir_count', 1),
                data.get('pir_price', 0)
            )
            self.increase_total_duration(doc, 3)

        # Шаг 3: Вычисляем итоги
        totals = self.calculate_totals(doc)

        # Шаг 4: Заменяем плейсхолдеры
        current_date = datetime.now().strftime("%d.%m.%Y")
        context = {
            'kp_name': data.get('kp_title', data.get('kp_name', '')),
            'date': current_date,
            'subtotal': self.format_price(totals.get('subtotal', 0)),
            'vat': self.format_price(totals.get('vat', 0)),
            'total_with_vat': self.format_price(totals.get('total_with_vat', 0)),
        }
        for key, value in totals.items():
            if key.startswith('stage') and key.endswith('_total'):
                context[key] = self.format_price(value)
        self.replace_placeholders(doc, context)

        # Шаг 4.1: Нормализуем границы таблицы (убираем пропуски чёрных линий)
        self.normalize_table_borders(doc)

        # Шаг 5: Сохранение
        # Очищаем запрещённые символы из имени файла
        safe_name = re.sub(r'[/\\:*?"<>|]', '_', data.get('kp_name', 'КП'))
        output_name = f"КП {safe_name}"
        save_dir = data.get('save_dir', os.path.dirname(os.path.abspath(__file__)))

        if not os.path.exists(save_dir):
            os.makedirs(save_dir)

        docx_path = os.path.normpath(os.path.join(save_dir, f"{output_name}.docx"))
        doc.save(docx_path)

        # Конвертация в PDF
        pdf_path = os.path.normpath(os.path.join(save_dir, f"{output_name}.pdf"))
        try:
            if os.name == 'nt':
                self._export_pdf_windows_word(docx_path, pdf_path)
            else:
                from docx2pdf import convert
                convert(docx_path, pdf_path)
        except Exception as e:
            print(f"Ошибка конвертации в PDF: {e}")
            if os.name == 'nt':
                try:
                    from docx2pdf import convert
                    convert(docx_path, pdf_path)
                except Exception:
                    pdf_path = None
            else:
                try:
                    libreoffice_bin = shutil.which("libreoffice") or shutil.which("soffice")
                    if not libreoffice_bin:
                        pdf_path = None
                    else:
                        subprocess.run(
                            [
                                libreoffice_bin,
                                "--headless",
                                "--convert-to",
                                "pdf",
                                docx_path,
                                "--outdir",
                                save_dir,
                            ],
                            stdout=subprocess.PIPE,
                            stderr=subprocess.PIPE,
                            timeout=45,
                            check=False,
                        )
                        if not os.path.exists(pdf_path):
                            pdf_path = None
                except Exception:
                    pdf_path = None

        return docx_path, pdf_path
