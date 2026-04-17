"""
contract_generator.py — генерация договора подряда в формате DOCX.

Исправления v3.1:
1. Дата в шапке — правое выравнивание
2. Форматирование наименований: 'Общество с ограниченной ответственностью' (строчные)
3. Нумерация раздела 3: 3.1, 3.2, 3.3
4. Сумма договора из КП (Итого с НДС 5%) + прописью
5. Аванс 30% от суммы + прописью
6. Склонение должности руководителя заказчика в родительный падеж
7. Склонение ФИО руководителя заказчика в родительный падеж
8. Таблица сметы копируется из КП без изменения форматирования
9. Приложение 3: горизонтальный лист при >12 месяцев, без зелёной заливки,
   корректное количество столбцов-месяцев
"""

import os
import re
import copy
import datetime
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import lxml.etree as etree


# ---------------------------------------------------------------------------
# Утилиты: числа прописью
# ---------------------------------------------------------------------------

def _num_to_words_ru(n: int) -> str:
    """Переводит целое неотрицательное число в текст на русском языке."""
    if n == 0:
        return 'ноль'

    units_m = ['', 'один', 'два', 'три', 'четыре', 'пять', 'шесть',
               'семь', 'восемь', 'девять', 'десять', 'одиннадцать',
               'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
               'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    units_f = ['', 'одна', 'две', 'три', 'четыре', 'пять', 'шесть',
               'семь', 'восемь', 'девять', 'десять', 'одиннадцать',
               'двенадцать', 'тринадцать', 'четырнадцать', 'пятнадцать',
               'шестнадцать', 'семнадцать', 'восемнадцать', 'девятнадцать']
    tens = ['', '', 'двадцать', 'тридцать', 'сорок', 'пятьдесят',
            'шестьдесят', 'семьдесят', 'восемьдесят', 'девяносто']
    hundreds = ['', 'сто', 'двести', 'триста', 'четыреста', 'пятьсот',
                'шестьсот', 'семьсот', 'восемьсот', 'девятьсот']

    def _three(num, feminine=False):
        parts = []
        h = num // 100
        t = (num % 100) // 10
        u = num % 10
        if h:
            parts.append(hundreds[h])
        if t == 1:
            parts.append(units_m[10 + u])
        else:
            if t:
                parts.append(tens[t])
            if u:
                parts.append(units_f[u] if feminine else units_m[u])
        return parts

    def _plural(num, one, two, five):
        if 11 <= num % 100 <= 19:
            return five
        r = num % 10
        if r == 1:
            return one
        if 2 <= r <= 4:
            return two
        return five

    result = []
    n = abs(int(n))

    billions = n // 1_000_000_000
    millions = (n % 1_000_000_000) // 1_000_000
    thousands = (n % 1_000_000) // 1_000
    remainder = n % 1_000

    if billions:
        result.extend(_three(billions))
        result.append(_plural(billions, 'миллиард', 'миллиарда', 'миллиардов'))
    if millions:
        result.extend(_three(millions))
        result.append(_plural(millions, 'миллион', 'миллиона', 'миллионов'))
    if thousands:
        result.extend(_three(thousands, feminine=True))
        result.append(_plural(thousands, 'тысяча', 'тысячи', 'тысяч'))
    if remainder:
        result.extend(_three(remainder))

    return ' '.join(result) if result else 'ноль'


def _amount_words(amount_str: str) -> str:
    """Возвращает сумму прописью (только рубли, без копеек)."""
    # Убираем пробелы, запятые и т.п.
    clean = re.sub(r'[\s\u00a0,]', '', str(amount_str))
    # Берём целую часть
    clean = clean.split('.')[0]
    try:
        n = int(clean)
        return _num_to_words_ru(n)
    except ValueError:
        return amount_str


def _format_amount(amount_str: str) -> str:
    """Форматирует число с пробелами-разделителями тысяч: '4644465' -> '4 644 465'."""
    clean = re.sub(r'[\s\u00a0,]', '', str(amount_str)).split('.')[0]
    try:
        n = int(clean)
        # Форматируем с пробелами
        s = f'{n:,}'.replace(',', '\u00a0')  # неразрывный пробел
        return s
    except ValueError:
        return amount_str


# ---------------------------------------------------------------------------
# Утилиты: склонение
# ---------------------------------------------------------------------------

TITLE_GENITIVE_MAP = {
    'директор': 'директора',
    'генеральный директор': 'генерального директора',
    'исполнительный директор': 'исполнительного директора',
    'технический директор': 'технического директора',
    'коммерческий директор': 'коммерческого директора',
    'финансовый директор': 'финансового директора',
    'президент': 'президента',
    'председатель': 'председателя',
    'управляющий': 'управляющего',
    'руководитель': 'руководителя',
    'индивидуальный предприниматель': 'индивидуального предпринимателя',
}

ADJ_GENITIVE_MAP = {
    'генеральный': 'генерального',
    'исполнительный': 'исполнительного',
    'технический': 'технического',
    'коммерческий': 'коммерческого',
    'финансовый': 'финансового',
    'операционный': 'операционного',
    'административный': 'административного',
    'региональный': 'регионального',
}

NOUN_GENITIVE_MAP = {
    'директор': 'директора',
    'президент': 'президента',
    'председатель': 'председателя',
    'управляющий': 'управляющего',
    'руководитель': 'руководителя',
    'предприниматель': 'предпринимателя',
    'начальник': 'начальника',
}


def _decline_title_genitive(title: str) -> str:
    """Склоняет должность в родительный падеж."""
    if not title:
        return title
    low = title.strip().lower()

    # Проверяем точные совпадения
    if low in TITLE_GENITIVE_MAP:
        return TITLE_GENITIVE_MAP[low]

    # Пословное склонение
    words = low.split()
    result = []
    for w in words:
        if w in ADJ_GENITIVE_MAP:
            result.append(ADJ_GENITIVE_MAP[w])
        elif w in NOUN_GENITIVE_MAP:
            result.append(NOUN_GENITIVE_MAP[w])
        elif w.endswith('ый') or w.endswith('ий'):
            result.append(w[:-2] + ('ого' if w.endswith('ый') else 'его'))
        elif w.endswith('ор') or w.endswith('ент') or w.endswith('ник'):
            result.append(w + 'а')
        elif w.endswith('ль'):
            result.append(w + 'я')
        else:
            result.append(w)
    return ' '.join(result)


def _decline_fio_genitive(fio: str) -> str:
    """Склоняет ФИО (мужской род) в родительный падеж."""
    if not fio:
        return fio
    parts = fio.strip().split()
    result = []
    for part in parts:
        if not part:
            continue
        low = part.lower()
        if low.endswith('ич'):          # Николаевич -> Николаевича
            result.append(part + 'а')
        elif low.endswith('ий'):        # Дмитрий -> Дмитрия
            result.append(part[:-2] + 'ия')
        elif low.endswith('ей'):        # Андрей -> Андрея
            result.append(part[:-2] + 'ея')
        elif low.endswith('ов') or low.endswith('ев') or low.endswith('ёв'):
            result.append(part + 'а')
        elif low.endswith('ин') or low.endswith('ын'):
            result.append(part + 'а')
        elif low.endswith('ай'):        # Николай -> Николая
            result.append(part[:-1] + 'я')
        elif low[-1] in 'бвгджзклмнпрстфхцчшщ':
            result.append(part + 'а')
        else:
            result.append(part)
    return ' '.join(result)


def _fio_to_short(fio: str) -> str:
    """
    Преобразует ФИО в формат Фамилия И.О.
    Примеры:
      'Иванов Иван Иванович' -> 'Иванов И.И.'
      'И.А. Сухов'       -> 'Сухов И.А.'  (инициалы перед фамилией)
    """
    if not fio:
        return fio
    fio = fio.strip()
    # Случай 1: уже в формате Фамилия И.О. — возвращаем как есть
    if re.match(r'^[A-ZА-Я][a-zа-я]+ [A-ZА-Я]\.[A-ZА-Я]\.$', fio):
        return fio
    # Случай 2: Формат 'И.А. Фамилия' (инициалы перед фамилией)
    m = re.match(r'^([A-ZА-Я]\.[A-ZА-Я]\.)\s+([A-ZА-Я][a-zа-я]+)$', fio)
    if m:
        initials, surname = m.group(1), m.group(2)
        return f'{surname} {initials}'
    # Случай 3: Формат 'Фамилия Имя Отчество'
    parts = fio.split()
    if len(parts) >= 3:
        return f'{parts[0]} {parts[1][0].upper()}.{parts[2][0].upper()}.'
    elif len(parts) == 2:
        return f'{parts[0]} {parts[1][0].upper()}.'
    else:
        return fio


# ---------------------------------------------------------------------------
# Утилиты: форматирование наименований
# ---------------------------------------------------------------------------

ORG_FORMS = {
    'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ': 'Общество с ограниченной ответственностью',
    'АКЦИОНЕРНОЕ ОБЩЕСТВО': 'Акционерное общество',
    'ПУБЛИЧНОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО': 'Публичное акционерное общество',
    'НЕПУБЛИЧНОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО': 'Непубличное акционерное общество',
    'ЗАКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО': 'Закрытое акционерное общество',
    'ОТКРЫТОЕ АКЦИОНЕРНОЕ ОБЩЕСТВО': 'Открытое акционерное общество',
    'ИНДИВИДУАЛЬНЫЙ ПРЕДПРИНИМАТЕЛЬ': 'Индивидуальный предприниматель',
    'МУНИЦИПАЛЬНОЕ УНИТАРНОЕ ПРЕДПРИЯТИЕ': 'Муниципальное унитарное предприятие',
    'ГОСУДАРСТВЕННОЕ УНИТАРНОЕ ПРЕДПРИЯТИЕ': 'Государственное унитарное предприятие',
    'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ УНИТАРНОЕ ПРЕДПРИЯТИЕ': 'Федеральное государственное унитарное предприятие',
    'НЕКОММЕРЧЕСКОЕ ПАРТНЁРСТВО': 'Некоммерческое партнёрство',
    'НЕКОММЕРЧЕСКОЕ ПАРТНЕРСТВО': 'Некоммерческое партнерство',
    'ТОВАРИЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ': 'Товарищество с ограниченной ответственностью',
    'ПРОИЗВОДСТВЕННЫЙ КООПЕРАТИВ': 'Производственный кооператив',
    'ПОТРЕБИТЕЛЬСКИЙ КООПЕРАТИВ': 'Потребительский кооператив',
    'ФОНД': 'Фонд',
    'УЧРЕЖДЕНИЕ': 'Учреждение',
    'АССОЦИАЦИЯ': 'Ассоциация',
    'СОЮЗ': 'Союз',
}


def _format_company_name_full(name: str) -> str:
    """
    Форматирует полное наименование организации из ВЕРХНЕГО РЕГИСТРА.
    Пример: 'ОБЩЕСТВО С ОГРАНИЧЕННОЙ ОТВЕТСТВЕННОСТЬЮ "НИКОЛЬСКИЙ ХЛЕБОЗАВОД"'
         -> 'Общество с ограниченной ответственностью «Никольский Хлебозавод»'
    """
    if not name:
        return name
    if name != name.upper():
        # Уже в нормальном регистре — только заменяем кавычки
        name = re.sub(r'"([^"]+)"', lambda m: '«' + m.group(1) + '»', name)
        # Исправляем написание ООО — если есть вариант с заглавными буквами
        name = re.sub(
            r'Общество С Ограниченной Ответственностью',
            'Общество с ограниченной ответственностью',
            name
        )
        return name

    # Ищем организационно-правовую форму
    name_upper = name.upper()
    org_form_result = ''
    rest = name

    for form_upper, form_normal in sorted(ORG_FORMS.items(), key=lambda x: -len(x[0])):
        if name_upper.startswith(form_upper):
            org_form_result = form_normal
            rest = name[len(form_upper):].strip()
            break

    # Обрабатываем остаток (название в кавычках)
    # Убираем ведущий пробел перед кавычкой
    rest = rest.strip()

    # Заменяем двойные кавычки на ёлочки и делаем Title Case внутри
    def replace_quotes(m):
        inner = m.group(1).strip()
        # Title Case для содержимого кавычек
        inner_titled = ' '.join(w.capitalize() for w in inner.split())
        return '«' + inner_titled + '»'

    rest_formatted = re.sub(r'"([^"]+)"', replace_quotes, rest)

    if org_form_result:
        if rest_formatted:
            return org_form_result + ' ' + rest_formatted
        return org_form_result
    else:
        # Не нашли форму — просто Title Case + замена кавычек
        return re.sub(r'"([^"]+)"', replace_quotes, name.capitalize())


def _format_company_name_short(name: str) -> str:
    """
    Форматирует краткое наименование.
    Пример: 'ООО "НИКОЛЬСКИЙ ХЛЕБОЗАВОД"' -> 'ООО «Никольский Хлебозавод»'
    """
    if not name:
        return name

    def replace_quotes(m):
        inner = m.group(1).strip()
        if inner == inner.upper() and any(c.isalpha() for c in inner):
            inner = ' '.join(w.capitalize() for w in inner.split())
        return '«' + inner + '»'

    result = re.sub(r'"([^"]+)"', replace_quotes, name)
    # Также обрабатываем уже существующие ёлочки с верхним регистром
    def replace_guillemets(m):
        inner = m.group(1).strip()
        if inner == inner.upper() and any(c.isalpha() for c in inner):
            inner = ' '.join(w.capitalize() for w in inner.split())
        return '«' + inner + '»'

    result = re.sub(r'«([^»]+)»', replace_guillemets, result)
    return result


# ---------------------------------------------------------------------------
# Утилиты: работа с DOCX
# ---------------------------------------------------------------------------

def _set_run_text(run, text: str):
    """Устанавливает текст run, сохраняя форматирование."""
    run.text = text


def _replace_paragraph_text(para, new_text: str):
    """Заменяет весь текст параграфа, сохраняя форматирование первого run."""
    # Удаляем гиперссылки (hyperlinks) из параграфа
    for hl in para._p.findall(qn('w:hyperlink')):
        para._p.remove(hl)
    if not para.runs:
        para.add_run(new_text)
        return
    # Сохраняем форматирование первого run
    first_run = para.runs[0]
    first_run.text = new_text
    # Удаляем остальные runs
    for run in para.runs[1:]:
        run.text = ''


def _copy_cell_format(src_tc, dst_tc):
    """Копирует форматирование ячейки (заливка, границы) из src в dst."""
    src_tcPr = src_tc.find(qn('w:tcPr'))
    if src_tcPr is None:
        return
    dst_tcPr = dst_tc.find(qn('w:tcPr'))
    if dst_tcPr is None:
        dst_tcPr = OxmlElement('w:tcPr')
        dst_tc.insert(0, dst_tcPr)
    # Копируем shd (заливку)
    src_shd = src_tcPr.find(qn('w:shd'))
    if src_shd is not None:
        dst_shd = dst_tcPr.find(qn('w:shd'))
        if dst_shd is not None:
            dst_tcPr.remove(dst_shd)
        dst_tcPr.append(copy.deepcopy(src_shd))


# ---------------------------------------------------------------------------
# Основной класс
# ---------------------------------------------------------------------------

class ContractGenerator:

    MONTHS_RU = [
        'Январь', 'Февраль', 'Март', 'Апрель', 'Май', 'Июнь',
        'Июль', 'Август', 'Сентябрь', 'Октябрь', 'Ноябрь', 'Декабрь'
    ]

    MONTHS_RU_GEN = [
        'января', 'февраля', 'марта', 'апреля', 'мая', 'июня',
        'июля', 'августа', 'сентября', 'октября', 'ноября', 'декабря'
    ]

    def __init__(self):
        base = os.path.dirname(os.path.abspath(__file__))
        self.template_path = os.path.join(base, 'templates', 'contract_template.docx')
        if not os.path.exists(self.template_path):
            raise FileNotFoundError(
                f'Шаблон договора не найден:\n{self.template_path}'
            )

    # -----------------------------------------------------------------------
    # Публичный метод
    # -----------------------------------------------------------------------

    def create_contract(self, data: dict) -> str:
        """
        Генерирует договор и возвращает путь к созданному файлу.

        data keys:
            contract_number, save_dir,
            customer_fullname, customer_shortname, customer_address,
            customer_ogrn, customer_inn, customer_kpp,
            customer_bank, customer_bik, customer_rs, customer_ks,
            customer_phone, customer_email,
            customer_director_title, customer_director_name, customer_basis,
            kp_file (optional)
        """
        doc = Document(self.template_path)
        today = datetime.date.today()

        # --- Извлекаем данные из КП ---
        kp_info = self._extract_kp_info(data.get('kp_file', ''))
        total_amount = kp_info.get('total_amount', '')
        months_count = kp_info.get('months_count', 12)
        kp_stages = kp_info.get('stages', [])

        # --- Форматируем наименования ---
        customer_fullname = _format_company_name_full(data.get('customer_fullname', ''))
        customer_shortname = _format_company_name_short(data.get('customer_shortname', ''))

        # --- Склоняем должность и ФИО ---
        dir_title_raw = data.get('customer_director_title', '')
        dir_title_gen = _decline_title_genitive(dir_title_raw)
        dir_name_raw = data.get('customer_director_name', '')
        dir_name_gen = _decline_fio_genitive(dir_name_raw)
        # Формат Фамилия И.О. для подписей в разделе 13 и приложениях
        dir_name_short = _fio_to_short(dir_name_raw)

        # --- Номер договора ---
        contract_num = data.get('contract_number', '')
        full_contract_num = f'0ЦЦБ-{contract_num}'
        contract_title = f'Договор подряда № {full_contract_num}'

        # --- Дата ---
        day = today.day
        month_gen = self.MONTHS_RU_GEN[today.month - 1]
        year = today.year
        date_str = f'«{day:02d}» {month_gen} {year} года.'
        date_str_short = f'«{day:02d}» {month_gen} {year} года'

        # --- Суммы ---
        total_val = 0.0
        if total_amount:
            clean = re.sub(r'[\s\u00a0]', '', str(total_amount)).replace(',', '.')
            try:
                total_val = float(clean)
            except ValueError:
                pass

        # Используем процент аванса из данных (по умолчанию 30%)
        try:
            advance_pct = float(str(data.get('advance_percent', '30')).replace(',', '.').strip())
            if advance_pct <= 0 or advance_pct >= 100:
                advance_pct = 30.0
        except (ValueError, TypeError):
            advance_pct = 30.0

        # Расчёт с копейками (округление до 2 знаков)
        advance_val = round(total_val * advance_pct / 100, 2) if total_val else 0.0
        postpay_val = round(total_val - advance_val, 2) if total_val else 0.0

        def _split_rub_kop(val):
            """Разделяет сумму на рубли (int) и копейки (int 0-99)."""
            rub = int(val)
            kop = round((val - rub) * 100)
            if kop >= 100:
                rub += 1
                kop = 0
            return rub, kop

        total_rub, total_kop = _split_rub_kop(total_val)
        advance_rub, advance_kop = _split_rub_kop(advance_val)
        postpay_rub, postpay_kop = _split_rub_kop(postpay_val)

        total_fmt = _format_amount(str(total_rub)) if total_rub else '****'
        total_words = _amount_words(str(total_rub)) if total_rub else '******'
        total_suffix = f' рублей, {total_kop:02d} копеек, с учетом НДС 5%.'
        advance_fmt = _format_amount(str(advance_rub)) if advance_rub else '****'
        advance_words = _amount_words(str(advance_rub)) if advance_rub else '******'
        advance_kop_str = f'{advance_kop:02d}'
        postpay_fmt = _format_amount(str(postpay_rub)) if postpay_rub else '****'
        postpay_words = _amount_words(str(postpay_rub)) if postpay_rub else '******'
        postpay_kop_str = f'{postpay_kop:02d}'

        # ===================================================================
        # Сохраняем ссылки на все нужные параграфы ДО вставки города
        # (вставка сдвигает индексы на +1)
        # ===================================================================
        p0 = doc.paragraphs[0]
        p2 = doc.paragraphs[2]
        p4_intro = doc.paragraphs[4]   # вводная часть
        p29_price = doc.paragraphs[29]  # цена договора 3.1/3.2
        p30_32 = doc.paragraphs[30]     # 3.3 -> 3.2
        p31_33 = doc.paragraphs[31]     # 3.4 -> 3.3
        p35_advance = doc.paragraphs[35]  # аванс 4.2

        # ===================================================================
        # 1. Заголовок договора P[0]
        # ===================================================================
        _replace_paragraph_text(p0, contract_title)

        # ===================================================================
        # 2. Дата P[2] — правое выравнивание, город убираем из строки даты
        # ===================================================================
        # Меняем выравнивание на правое
        pPr = p2._p.find(qn('w:pPr'))
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            p2._p.insert(0, pPr)
        jc = pPr.find(qn('w:jc'))
        if jc is None:
            jc = OxmlElement('w:jc')
            pPr.append(jc)
        jc.set(qn('w:val'), 'right')

        # Извлекаем название города из первого run и убираем его из строки даты
        city_name = 'г. Иркутск'
        if p2.runs:
            city_run = p2.runs[0]
            city_text = city_run.text.strip()
            # Если в первом run есть название города — сохраняем его
            if city_text and ('г.' in city_text or 'город' in city_text.lower()):
                city_name = city_text
            # Убираем город из строки даты
            city_run.text = ''
            # Очищаем runs 1 и 2 (лишние пробелы)
            for run in p2.runs[1:3]:
                run.text = ''
            # run[3] = '«'
            if len(p2.runs) > 3:
                p2.runs[3].text = '«'
            # run[4] = день
            if len(p2.runs) > 4:
                p2.runs[4].text = f'{day:02d}'
            # run[5] = '»'
            if len(p2.runs) > 5:
                p2.runs[5].text = '»'
            # run[6] = ' '
            if len(p2.runs) > 6:
                p2.runs[6].text = ' '
            # run[7] = месяц
            if len(p2.runs) > 7:
                p2.runs[7].text = month_gen
            # run[8] = ' YYYY' (год)
            if len(p2.runs) > 8:
                p2.runs[8].text = f' {year}'
            # run[9] = '' (очищаем)
            if len(p2.runs) > 9:
                p2.runs[9].text = ''
            # run[10] = ' г'
            if len(p2.runs) > 10:
                p2.runs[10].text = ' г'
            # run[11] = 'ода'
            if len(p2.runs) > 11:
                p2.runs[11].text = 'ода'
            # run[12] = '.'
            if len(p2.runs) > 12:
                p2.runs[12].text = '.'

        # Добавляем параграф с городом после строки даты (справа)
        # Находим параграф P[2] в XML и вставляем после него новый параграф с городом
        city_para = OxmlElement('w:p')
        city_pPr = OxmlElement('w:pPr')
        city_jc = OxmlElement('w:jc')
        city_jc.set(qn('w:val'), 'right')
        city_pPr.append(city_jc)
        # Копируем стиль шрифта из P[2]
        if p2.runs:
            rPr_src = p2.runs[0]._r.find(qn('w:rPr'))
            if rPr_src is not None:
                city_rPr = copy.deepcopy(rPr_src)
            else:
                city_rPr = None
        else:
            city_rPr = None
        city_para.append(city_pPr)
        city_r = OxmlElement('w:r')
        if city_rPr is not None:
            city_r.append(copy.deepcopy(city_rPr))
        city_t = OxmlElement('w:t')
        city_t.text = city_name
        city_r.append(city_t)
        city_para.append(city_r)
        # Вставляем после P[2]
        p2._p.addnext(city_para)

        # ===================================================================
        # 3. Вводная часть — заказчик (используем сохранённую ссылку)
        # ===================================================================
        p4 = p4_intro
        basis = data.get('customer_basis', 'Устава')
        # Обновляем runs P[4] точечно по индексам шаблона:
        # run[0-4] = ООО «Гидросервис» (bold) - не меняем
        # run[5] = ',' - не меняем
        # run[6] = ' именуемое в дальнейшем ' - не меняем
        # run[7] = '«Подрядчик»' (bold) - не меняем
        # run[8] = ', в лице Генерального директора Сухова...' - оставляем как есть
        # run[9] = 'с одной стороны ' - оставляем
        # run[10] = 'и, ' - оставляем
        # run[11] = 'Государственное б' (bold) -> наименование заказчика
        # run[12-13] = продолжение наименования (bold) -> очистить
        # run[14] = '(ГБУЗ МО «Санаторий Пушкино»)' (bold) -> краткое наименование
        # run[15] = ',' (bold) -> оставляем
        # run[16] = ' ' (bold) -> оставляем
        # run[17] = 'именуемое в дальнейшем ' - не меняем
        # run[18] = '«Заказчик»' (bold) - не меняем
        # run[19] = ',' - не меняем
        # run[20] = ' ' - не меняем
        # run[21] = 'в лице ' - не меняем
        # run[22] = 'д' -> должность заказчика
        # run[23] = 'иректора ' -> очистить
        # run[24] = 'Хромовских Игоря Валерьевича' -> ФИО заказчика
        # run[25] = ', ' - не меняем
        # run[26] = 'действующего на основании ' - не меняем
        # run[27] = 'Устава' -> основание
        # run[28] = ' ' - не меняем
        # run[29] = 'с другой стороны...' - не меняем
        runs4 = p4.runs
        # Наименование заказчика (bold): run[11] = полное, run[12-13] = очистить, run[14] = краткое
        if len(runs4) > 11:
            runs4[11].text = customer_fullname
        if len(runs4) > 12:
            runs4[12].text = ''
        if len(runs4) > 13:
            runs4[13].text = ''
        if len(runs4) > 14:
            runs4[14].text = f' ({customer_shortname})'
        # Должность заказчика: run[22] = должность, run[23] = очистить
        if len(runs4) > 22:
            runs4[22].text = dir_title_gen + ' '
        if len(runs4) > 23:
            runs4[23].text = ''
        # ФИО заказчика: run[24]
        if len(runs4) > 24:
            runs4[24].text = dir_name_gen
        # Основание: run[27]
        if len(runs4) > 27:
            runs4[27].text = basis

        # ===================================================================
        # 3.1. Пункт 1.5 (адрес проведения работ)
        # ===================================================================
        if data.get('include_work_address'):
            self._insert_work_address_clause(doc, data.get('work_address', ''))

        # ===================================================================
        # 4. Раздел 3 — нумерация и суммы
        # ===================================================================
        # P[29]: 3.2 -> 3.1 + сумма из КП
        self._update_price_paragraph(
            p29_price,
            prefix='3.1. ',
            amount_fmt=total_fmt,
            amount_words=total_words,
            suffix=total_suffix
        )

        # P[30]: 3.3 -> 3.2
        text30 = p30_32.text
        if text30.startswith('3.3'):
            _replace_paragraph_text(p30_32, '3.2' + text30[3:])

        # P[31]: 3.4 -> 3.3
        text31 = p31_33.text
        if text31.startswith('3.4'):
            _replace_paragraph_text(p31_33, '3.3' + text31[3:])

        # ===================================================================
        # 5. Раздел 4 — аванс P[35]
        # ===================================================================
        self._update_advance_paragraph(
            p35_advance,
            advance_fmt=advance_fmt,
            advance_words=advance_words,
            postpay_fmt=postpay_fmt,
            postpay_words=postpay_words,
            advance_kop=advance_kop_str,
            postpay_kop=postpay_kop_str
        )

        # ===================================================================
        # 6. Реквизиты сторон (таблица 0)
        # ===================================================================
        self._fill_requisites_table(doc, data, customer_fullname, customer_shortname,
                                    dir_title_raw, dir_name_raw)

        # Удаляем лишние пустые параграфы между таблицей реквизитов и Приложением 1
        # (в шаблоне 4 пустых параграфа, оставляем 1)
        body = doc.element.body
        tbl0_elem = doc.tables[0]._tbl
        sibling = tbl0_elem.getnext()
        empty_paras = []
        while sibling is not None:
            stag = sibling.tag.split('}')[-1] if '}' in sibling.tag else sibling.tag
            if stag == 'p':
                from docx.text.paragraph import Paragraph as _Para
                pp = _Para(sibling, doc)
                if not pp.text.strip():
                    empty_paras.append(sibling)
                else:
                    break  # нашли непустой параграф (Приложение 1)
            else:
                break
            sibling = sibling.getnext()
        # Удаляем все кроме первого пустого параграфа
        for ep in empty_paras[1:]:
            body.remove(ep)

        # ===================================================================
        # 7. Заголовки приложений
        # ===================================================================
        self._update_appendix_headers(doc, full_contract_num, date_str_short)

        # ===================================================================
        # 7а. Подписи в разделе 13 и приложениях (Фамилия И.О.)
        # ===================================================================
        self._update_signatures(doc, dir_title_raw, dir_name_raw, customer_shortname)

        # ===================================================================
        # 8. Смета — Приложение №1 (таблица 1)
        # ===================================================================
        kp_file = data.get('kp_file', '')
        if kp_file and os.path.exists(kp_file):
            self._replace_estimate_table(doc, kp_file)
            self._add_vat_column_to_estimate(doc)

        # ===================================================================
        # 9. График работ — Приложение №3 (таблица 4)
        #    + горизонтальный лист если >12 месяцев
        # ===================================================================
        self._build_schedule_table(doc, kp_stages, months_count, today)
        if months_count > 12:
            self._set_landscape_for_app3(doc)

        # ===================================================================
        # 10. Сохранение
        # ===================================================================
        save_dir = data.get('save_dir', os.path.expanduser('~'))
        os.makedirs(save_dir, exist_ok=True)
        filename = f'Договор подряда № {full_contract_num}.docx'
        out_path = os.path.join(save_dir, filename)
        doc.save(out_path)
        return out_path

    # -----------------------------------------------------------------------
    # Вспомогательные методы
    # -----------------------------------------------------------------------

    def _extract_kp_info(self, kp_file: str) -> dict:
        """Извлекает данные из файла КП: сумму, количество месяцев, этапы."""
        result = {'total_amount': '', 'months_count': 12, 'stages': []}
        if not kp_file or not os.path.exists(kp_file):
            return result

        try:
            kp_doc = Document(kp_file)
            if not kp_doc.tables:
                return result

            t = kp_doc.tables[0]
            stages = []
            months_count = 12

            for ri, row in enumerate(t.rows):
                cells = [c.text.strip() for c in row.cells]
                if len(cells) < 2:
                    continue

                # Ищем строку с "Итого с НДС 5%"
                row_text = ' '.join(cells).lower()
                if 'итого с ндс' in row_text or 'итого с нд' in row_text:
                    # Берём последнюю непустую ячейку как сумму
                    for c in reversed(cells):
                        c_clean = c.strip()
                        if c_clean:
                            result['total_amount'] = c_clean
                            break

                # Извлекаем количество месяцев из строки 0, столбец 2
                if ri == 0 and len(cells) >= 3:
                    cell_val = cells[2]
                    # Ищем число перед словом 'мес'
                    m = re.search(r'(\d+)\s*мес', cell_val, re.IGNORECASE)
                    if m:
                        months_count = int(m.group(1))
                    else:
                        # Ищем любое число
                        m2 = re.search(r'(\d+)', cell_val)
                        if m2:
                            months_count = int(m2.group(1))

                # Собираем ВСЕ строки сметы (этапы и подэтапы: 1, 1.1, 1.2, 2, 2.1 и т.д.)
                # Исключаем: заголовок (ri==0), итоговые строки (итого), пустые строки
                num_cell = cells[0].strip()
                name_cell = cells[1].strip() if len(cells) > 1 else ''
                # Пропускаем заголовок, итоговые строки и строки без номера
                if ri == 0:
                    continue  # заголовок
                if not num_cell or not name_cell:
                    continue  # пустая строка
                if 'итого' in name_cell.lower() or 'итого' in num_cell.lower():
                    continue  # итоговая строка
                # Принимаем строки с номером вида: 1, 2, 3, 1.1, 1.2, 2.1 и т.д.
                if re.match(r'^\d+(\.\d+)*$', num_cell):
                    stages.append({'num': num_cell, 'name': name_cell})

            result['months_count'] = months_count
            result['stages'] = stages

        except Exception as e:
            print(f'Ошибка извлечения данных из КП: {e}')

        return result

    def _update_price_paragraph(self, para, prefix: str, amount_fmt: str,
                                 amount_words: str, suffix: str):
        """Обновляет параграф с ценой договора."""
        # Строим новый текст с нужными runs
        # Структура: prefix + amount_fmt (bold) + ' (' + amount_words + ')' + suffix
        for run in para.runs:
            run.text = ''

        if para.runs:
            para.runs[0].text = prefix + 'Цена работ по настоящему Договору составляет '
            if len(para.runs) > 1:
                para.runs[1].text = amount_fmt
                para.runs[1].bold = True
            if len(para.runs) > 2:
                para.runs[2].text = ' ('
            if len(para.runs) > 3:
                para.runs[3].text = amount_words
            if len(para.runs) > 4:
                para.runs[4].text = ')'
            if len(para.runs) > 5:
                para.runs[5].text = suffix
        else:
            r = para.add_run(prefix + 'Цена работ по настоящему Договору составляет ')
            r2 = para.add_run(amount_fmt)
            r2.bold = True
            para.add_run(f' ({amount_words}){suffix}')

    def _update_advance_paragraph(self, para, advance_fmt: str, advance_words: str,
                                   postpay_fmt: str, postpay_words: str,
                                   advance_kop: str = '00', postpay_kop: str = '00'):
        """
        Точечно обновляет параграф 4.2 (аванс) по фиксированным индексам runs.
        Структура шаблона:
          run[5]  = сумма аванса (bold)
          run[7]  = слова аванса часть 1
          run[8]  = слова аванса часть 2 (очистить)
          run[9]  = пробел (очистить)
          run[10] = слова аванса часть 3 (очистить)
          run[16] = копейки аванса (bold)
          run[29] = сумма постоплаты часть 1 (bold)
          run[30] = сумма постоплаты часть 2 (bold, очистить)
          run[32] = слова постоплаты часть 1
          run[33] = пробел (очистить)
          run[34] = слова постоплаты часть 2 + ')' -> оставить только ') '
          run[37] = копейки постоплаты (bold)
        """
        runs = para.runs
        if not runs:
            return

        # Аванс: сумма
        if len(runs) > 5:
            runs[5].text = advance_fmt
        # Аванс: пропись (в первый run, остальные очистить)
        if len(runs) > 7:
            runs[7].text = advance_words
        if len(runs) > 8:
            runs[8].text = ''
        if len(runs) > 9:
            runs[9].text = ''
        if len(runs) > 10:
            runs[10].text = ''
        # Аванс: копейки
        if len(runs) > 16:
            runs[16].text = advance_kop

        # Постоплата: сумма
        if len(runs) > 29:
            runs[29].text = postpay_fmt
        if len(runs) > 30:
            runs[30].text = ''
        # Постоплата: пропись
        if len(runs) > 32:
            runs[32].text = postpay_words
        if len(runs) > 33:
            runs[33].text = ''
        if len(runs) > 34:
            runs[34].text = ') '
        # Постоплата: копейки
        if len(runs) > 37:
            runs[37].text = postpay_kop

    def _insert_work_address_clause(self, doc, work_address: str):
        """
        Добавляет в раздел 1 пункт:
        1.5.  Работы проводятся по адресу: "<адрес>"
        """
        address = (work_address or '').strip()
        if not address:
            return

        target_para = None
        for para in doc.paragraphs:
            if re.match(r'^\s*1\.4\.', para.text.strip()):
                target_para = para
                break

        # Запасной вариант: вставляем после последнего пункта раздела 1
        if target_para is None:
            for para in doc.paragraphs:
                if re.match(r'^\s*1\.\d+\.', para.text.strip()):
                    target_para = para

        if target_para is None:
            return

        new_p = copy.deepcopy(target_para._p)
        target_para._p.addnext(new_p)
        from docx.text.paragraph import Paragraph as _Paragraph
        new_para = _Paragraph(new_p, doc)
        _replace_paragraph_text(
            new_para,
            f'1.5.  Работы проводятся по адресу: "{address}"'
        )

    def _fill_requisites_table(self, doc, data: dict, customer_fullname: str,
                                customer_shortname: str, dir_title_raw: str,
                                dir_name_raw: str):
        """Заполняет таблицу реквизитов сторон (таблица 0)."""
        t = doc.tables[0]
        # Таблица имеет 4 строки:
        # Строка 0: заголовки «ПОДРЯДЧИК» / «ЗАКАЗЧИК»
        # Строка 1: наименования
        # Строка 2: реквизиты
        # Строка 3: подписи

        # Строка 1, ячейка 1 — наименование заказчика
        if len(t.rows) > 1 and len(t.rows[1].cells) > 1:
            cell = t.rows[1].cells[1]
            paras_r1 = cell.paragraphs
            name_set = False
            for p in paras_r1:
                txt = p.text.strip()
                if not name_set and txt:
                    # Первый непустой параграф — сокращённое название организации
                    _replace_paragraph_text(p, customer_shortname)
                    name_set = True
                else:
                    # Все остальные параграфы (в т.ч. с /ФИО/) — очищаем
                    _replace_paragraph_text(p, '')

        # Строка 2, ячейка 1 — реквизиты заказчика
        if len(t.rows) > 2 and len(t.rows[2].cells) > 1:
            cell = t.rows[2].cells[1]
            addr = data.get('customer_address', '')
            ogrn = data.get('customer_ogrn', '')
            inn = data.get('customer_inn', '')
            kpp = data.get('customer_kpp', '')
            bank = data.get('customer_bank', '')
            bik = data.get('customer_bik', '')
            rs = data.get('customer_rs', '')
            ks_raw = data.get('customer_ks', '')
            # Очищаем корр. счёт если в поле попал JSON-объект (старый баг API)
            if isinstance(ks_raw, dict):
                ks = ks_raw.get('Номер', '') or ''
            elif isinstance(ks_raw, str) and ks_raw.strip().startswith('{'):
                # Строка вида "{'\u041dомер': '30101...', '\u0414ата': '...'}"
                import ast
                try:
                    parsed = ast.literal_eval(ks_raw)
                    ks = parsed.get('Номер', '') if isinstance(parsed, dict) else ks_raw
                except Exception:
                    # Если не удалось разобрать — извлекаем числа регекспом
                    nums = re.findall(r'\d{15,}', ks_raw)
                    ks = nums[0] if nums else ''
            else:
                ks = ks_raw
            phone = data.get('customer_phone', '')
            email = data.get('customer_email', '')

            lines = []
            if addr:
                lines.append(f'Юр./Факт./Почт. адрес: {addr}')
            if ogrn:
                lines.append(f'ОГРН: {ogrn}')
            if inn:
                lines.append(f'ИНН: {inn}')
            if kpp:
                lines.append(f'КПП: {kpp}')
            if bank:
                lines.append(f'Банк: {bank}')
            if bik:
                lines.append(f'БИК: {bik}')
            if rs:
                lines.append(f'Р/с: {rs}')
            if ks:
                lines.append(f'К/с: {ks}')
            if phone:
                lines.append(f'Тел.: {phone}')
            if email:
                lines.append(f'E-mail: {email}')

            # Заполняем параграфы ячейки
            paras = cell.paragraphs
            for i, line in enumerate(lines):
                if i < len(paras):
                    _replace_paragraph_text(paras[i], line)
                else:
                    cell.add_paragraph(line)
            # Очищаем лишние параграфы
            for i in range(len(lines), len(paras)):
                _replace_paragraph_text(paras[i], '')

        # Строка 3, ячейка 0 — подпись подрядчика (Гидросервис)
        # Приводим /И.А. Сухов/ к формату /Сухов И.А./
        if len(t.rows) > 3 and len(t.rows[3].cells) > 0:
            cell0 = t.rows[3].cells[0]
            for p in cell0.paragraphs:
                txt = p.text.strip()
                if '/' in txt and txt.count('/') >= 2 and 'Сухов' in txt:
                    # Извлекаем ФИО между слешами: "____ /И.А. Сухов/" -> "И.А. Сухов"
                    m = re.search(r'/([^/]+)/', txt)
                    if m:
                        fio_inner = m.group(1).strip()
                        normalized = _fio_to_short(fio_inner)
                        # Сохраняем подчёркивания перед слешами
                        prefix = txt[:txt.index('/')]
                        _replace_paragraph_text(p, f'{prefix}/{normalized}/')
                    break

        # Строка 3, ячейка 1 — подпись заказчика
        if len(t.rows) > 3 and len(t.rows[3].cells) > 1:
            cell = t.rows[3].cells[1]
            paras = cell.paragraphs
            # Ищем параграф с должностью
            for p in paras:
                text = p.text.strip()
                if 'директор' in text.lower() or 'руководитель' in text.lower() or 'директора' in text.lower():
                    _replace_paragraph_text(p, f'{dir_title_raw.capitalize()}:')
                    break
            # Ищем параграф с ФИО — используем формат Фамилия И.О.
            dir_name_short = _fio_to_short(dir_name_raw)
            for p in paras:
                text = p.text.strip()
                if '/' in text and text.count('/') >= 2:
                    # Сохраняем подчёркивания перед слешами
                    prefix = text[:text.index('/')]
                    _replace_paragraph_text(p, f'{prefix}/{dir_name_short}/')
                    break

    def _update_appendix_headers(self, doc, full_contract_num: str, date_str: str):
        """Обновляет заголовки приложений (номер и дата договора)."""
        for i, p in enumerate(doc.paragraphs):
            text = p.text.strip()
            # Строки вида "к Договору подряда № 0ЦЦБ-..."
            if text.startswith('к Договору подряда №'):
                _replace_paragraph_text(p, f'к Договору подряда № {full_contract_num}')
            # Строки вида "от «02» февраля 2026 года."
            elif text.startswith('от «') and 'года' in text:
                _replace_paragraph_text(p, f'от {date_str}')

    def _replace_estimate_table(self, doc, kp_file: str):
        """Заменяет таблицу сметы в договоре на таблицу из КП (полная копия)."""
        try:
            kp_doc = Document(kp_file)
            if not kp_doc.tables:
                return

            kp_tbl = kp_doc.tables[0]._tbl
            contract_tbl = doc.tables[1]._tbl

            # Заменяем tbl элемент в body
            parent = contract_tbl.getparent()
            idx = list(parent).index(contract_tbl)
            new_tbl = copy.deepcopy(kp_tbl)
            parent.remove(contract_tbl)
            parent.insert(idx, new_tbl)

        except Exception as e:
            print(f'Ошибка замены таблицы сметы: {e}')

    def _add_vat_column_to_estimate(self, doc):
        """
        Добавляет столбец "Цена с НДС 5%" в таблицу сметы (Приложение Ⅶ1).
        Столбец добавляется после последнего столбца с ценой (без НДС).
        Для каждой строки вычисляет цену с НДС = цена без НДС * 1.05.
        """
        try:
            # После замены таблица сметы стала таблицей[1]
            estimate_tbl = doc.tables[1]
            tbl_el = estimate_tbl._tbl

            # Определяем индекс столбца с ценой (без НДС) в заголовке
            # Ищем в первых 2 строках ячейку с "цена" или "стоимость" (без НДС)
            price_col_idx = -1
            rows = estimate_tbl.rows
            if not rows:
                return

            # Ищем столбец с ценой в заголовке (1-2 строка)
            for ri in range(min(2, len(rows))):
                cells = rows[ri].cells
                for ci, cell in enumerate(cells):
                    ct = cell.text.strip().lower()
                    if ('цена' in ct or 'стоимость' in ct) and 'ндс' not in ct:
                        price_col_idx = ci
                        break
                if price_col_idx >= 0:
                    break

            # Если не нашли — берём последний столбец
            if price_col_idx < 0:
                price_col_idx = len(rows[0].cells) - 1

            # Обрабатываем каждую строку таблицы
            total_rows = len(rows)
            for ri, row in enumerate(rows):
                cells = row.cells
                if not cells:
                    continue

                # Шаблонная ячейка для копирования форматирования
                template_tc = cells[price_col_idx]._tc

                # Создаём новую ячейку на основе шаблонной
                new_tc = copy.deepcopy(template_tc)

                # Определяем текст новой ячейки
                if ri == 0:
                    # Заголовок столбца
                    cell_text = 'Цена с НДС 5%, руб.'
                elif ri >= total_rows - 2:
                    # 2 последних строки (итоговые суммы) — оставляем пустыми
                    cell_text = ''
                else:
                    # Для строк данных — вычисляем цену с НДС
                    price_text = cells[price_col_idx].text.strip()
                    # Очищаем число от пробелов, запятых и неразрывных пробелов
                    price_clean = re.sub(r'[\s\u00a0 ]', '', price_text).replace(',', '.')
                    try:
                        price_val = float(price_clean)
                        price_with_vat = price_val * 1.05
                        # Форматируем: если целое — без дроби, если дробь — с 2 знаками
                        if price_with_vat == int(price_with_vat):
                            cell_text = _format_amount(str(int(price_with_vat)))
                        else:
                            cell_text = f'{price_with_vat:,.2f}'.replace(',', ' ')
                    except (ValueError, TypeError):
                        # Если не число — оставляем пустым
                        cell_text = ''

                # Записываем текст в новую ячейку
                self._set_tc_text(new_tc, cell_text)

                # Добавляем ячейку в конец строки
                row._tr.append(new_tc)

            # Обновляем gridCol в tblGrid для нового столбца
            tblGrid = tbl_el.find(qn('w:tblGrid'))
            if tblGrid is not None:
                # Копируем последний gridCol
                grid_cols = tblGrid.findall(qn('w:gridCol'))
                if grid_cols:
                    new_grid_col = copy.deepcopy(grid_cols[-1])
                    tblGrid.append(new_grid_col)

            # Уменьшаем ширину всех столбцов для вмещаемости
            # Читаем текущую ширину таблицы
            tblPr = tbl_el.find(qn('w:tblPr'))
            if tblPr is not None:
                tblW = tblPr.find(qn('w:tblW'))
                if tblW is not None:
                    # Устанавливаем ширину таблицы = 100% страницы
                    tblW.set(qn('w:w'), '5000')
                    tblW.set(qn('w:type'), 'pct')

        except Exception as e:
            print(f'Ошибка добавления столбца НДС: {e}')
            import traceback
            traceback.print_exc()

    def _build_schedule_table(self, doc, stages: list, months_count: int,
                               today: datetime.date):
        """
        Строит таблицу графика работ в Приложении №3.
        Строки = этапы из сметы, столбцы = месяцы (начиная с текущего).
        """
        try:
            # Таблица 4 в шаблоне — это таблица с индексом 4
            # Но после замены таблицы 1 индексы могут сдвинуться
            # Ищем таблицу по содержимому заголовка
            schedule_tbl = None
            for t in doc.tables:
                # Ищем таблицу с "Наименование Работ" в заголовке
                for row in t.rows[:2]:
                    for cell in row.cells:
                        if 'Наименование' in cell.text or 'наименование' in cell.text:
                            schedule_tbl = t
                            break
                    if schedule_tbl:
                        break
                if schedule_tbl:
                    break

            if schedule_tbl is None:
                return

            tbl_element = schedule_tbl._tbl
            rows = tbl_element.findall(qn('w:tr'))

            if len(rows) < 2:
                return

            # Строка 0 — заголовок с годами
            # Строка 1 — заголовок с месяцами
            # Строка 2+ — данные (шаблонная строка данных)

            header_row_0 = rows[0]  # строка с годами
            header_row_1 = rows[1]  # строка с месяцами
            template_data_row = rows[2] if len(rows) > 2 else None  # шаблонная строка данных

            # --- Обновляем строку 0 (годы) ---
            self._update_year_header_row(header_row_0, months_count, today)

            # --- Обновляем строку 1 (месяцы) ---
            self._update_month_header_row(header_row_1, months_count, today)

            # --- Удаляем все строки данных (строки 2+) ---
            for row in rows[2:]:
                tbl_element.remove(row)

            # --- Добавляем строки для каждого этапа ---
            if not stages:
                # Если нет этапов — добавляем одну пустую строку
                if template_data_row is not None:
                    new_row = copy.deepcopy(template_data_row)
                    self._clear_row_green(new_row)
                    # Добавляем нужное количество ячеек для месяцев
                    self._ensure_row_cells(new_row, months_count + 2)
                    tbl_element.append(new_row)
            else:
                for stage in stages:
                    if template_data_row is not None:
                        new_row = copy.deepcopy(template_data_row)
                    else:
                        new_row = copy.deepcopy(header_row_1)

                    self._clear_row_green(new_row)
                    self._ensure_row_cells(new_row, months_count + 2)

                    # Заполняем номер и наименование
                    tcs = new_row.findall(qn('w:tc'))
                    if len(tcs) >= 1:
                        self._set_tc_text(tcs[0], stage['num'])
                    if len(tcs) >= 2:
                        self._set_tc_text(tcs[1], stage['name'])
                    # Остальные ячейки — пустые
                    for tc in tcs[2:]:
                        self._set_tc_text(tc, '')

                    tbl_element.append(new_row)

            # --- Выравниваем ширину столбцов месяцев ---
            # Альбомная страница: ширина ~15840 twips (11 дюймов)
            # Поля: по 720 twips (0.5") с каждой стороны
            # Доступная ширина: 15840 - 720*2 = 14400 twips
            total_width = 14400
            num_col_width = 700     # столбец №
            name_col_width = 3500   # столбец Наименование
            remaining = total_width - num_col_width - name_col_width
            month_col_width = remaining // months_count if months_count > 0 else 700

            # Обновляем tblGrid
            tblGrid = tbl_element.find(qn('w:tblGrid'))
            if tblGrid is not None:
                # Удаляем старые gridCol
                for gc in tblGrid.findall(qn('w:gridCol')):
                    tblGrid.remove(gc)
                # Добавляем новые: №, Наименование, месяцы
                gc_num = OxmlElement('w:gridCol')
                gc_num.set(qn('w:w'), str(num_col_width))
                tblGrid.append(gc_num)
                gc_name = OxmlElement('w:gridCol')
                gc_name.set(qn('w:w'), str(name_col_width))
                tblGrid.append(gc_name)
                for _ in range(months_count):
                    gc_m = OxmlElement('w:gridCol')
                    gc_m.set(qn('w:w'), str(month_col_width))
                    tblGrid.append(gc_m)

            # Устанавливаем ширину ячеек во всех строках (кроме строки годов с gridSpan)
            all_rows = tbl_element.findall(qn('w:tr'))
            for r_idx, tr in enumerate(all_rows):
                tcs = tr.findall(qn('w:tc'))
                for c_idx, tc in enumerate(tcs):
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    # Пропускаем ячейки с gridSpan (строка годов)
                    gs = tcPr.find(qn('w:gridSpan'))
                    if gs is not None:
                        span = int(gs.get(qn('w:val'), '1'))
                        if c_idx == 0:
                            w_val = num_col_width
                        elif c_idx == 1:
                            w_val = name_col_width
                        else:
                            w_val = month_col_width * span
                    else:
                        if c_idx == 0:
                            w_val = num_col_width
                        elif c_idx == 1:
                            w_val = name_col_width
                        else:
                            w_val = month_col_width
                    tcW = tcPr.find(qn('w:tcW'))
                    if tcW is None:
                        tcW = OxmlElement('w:tcW')
                        tcPr.append(tcW)
                    tcW.set(qn('w:w'), str(w_val))
                    tcW.set(qn('w:type'), 'dxa')

        except Exception as e:
            print(f'Ошибка построения графика работ: {e}')
            import traceback
            traceback.print_exc()

    def _update_year_header_row(self, header_row, months_count: int, today: datetime.date):
        """Обновляет строку заголовка с годами."""
        tcs = header_row.findall(qn('w:tc'))
        if len(tcs) < 3:
            return

        # Определяем распределение месяцев по годам
        year_months = {}  # {year: count}
        for i in range(months_count):
            month_idx = (today.month - 1 + i) % 12
            year_offset = (today.month - 1 + i) // 12
            yr = today.year + year_offset
            year_months[yr] = year_months.get(yr, 0) + 1

        # Ячейки 0 и 1 — фиксированные (№, Наименование)
        # Ячейки 2+ — годы с gridSpan

        # Удаляем все ячейки начиная с индекса 2
        for tc in tcs[2:]:
            header_row.remove(tc)

        # Добавляем ячейки для каждого года
        template_tc = tcs[2] if len(tcs) > 2 else copy.deepcopy(tcs[1])

        for yr, cnt in sorted(year_months.items()):
            new_tc = copy.deepcopy(template_tc if len(tcs) > 2 else tcs[-1])
            self._clear_row_green_tc(new_tc)
            # Устанавливаем gridSpan
            tcPr = new_tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                new_tc.insert(0, tcPr)
            gridSpan = tcPr.find(qn('w:gridSpan'))
            if gridSpan is None:
                gridSpan = OxmlElement('w:gridSpan')
                tcPr.append(gridSpan)
            gridSpan.set(qn('w:val'), str(cnt))
            # Устанавливаем текст года
            self._set_tc_text(new_tc, str(yr))
            # Центрирование
            self._set_tc_align(new_tc, 'center')
            header_row.append(new_tc)

    def _update_month_header_row(self, header_row, months_count: int, today: datetime.date):
        """Обновляет строку заголовка с названиями месяцев."""
        tcs = header_row.findall(qn('w:tc'))
        if len(tcs) < 3:
            return

        # Шаблонная ячейка месяца
        template_month_tc = tcs[2] if len(tcs) > 2 else copy.deepcopy(tcs[-1])

        # Удаляем все ячейки начиная с индекса 2
        for tc in tcs[2:]:
            header_row.remove(tc)

        # Добавляем ячейки для каждого месяца
        for i in range(months_count):
            month_idx = (today.month - 1 + i) % 12
            month_name = self.MONTHS_RU[month_idx]

            new_tc = copy.deepcopy(template_month_tc)
            self._clear_row_green_tc(new_tc)

            # Убираем gridSpan если есть
            tcPr = new_tc.find(qn('w:tcPr'))
            if tcPr is not None:
                gridSpan = tcPr.find(qn('w:gridSpan'))
                if gridSpan is not None:
                    tcPr.remove(gridSpan)

            self._set_tc_text(new_tc, month_name)
            self._set_tc_align(new_tc, 'center')
            header_row.append(new_tc)

    def _ensure_row_cells(self, row, target_count: int):
        """Обеспечивает нужное количество ячеек в строке данных."""
        tcs = row.findall(qn('w:tc'))
        current = len(tcs)

        if current < target_count:
            # Добавляем ячейки
            template_tc = tcs[-1] if tcs else None
            for _ in range(target_count - current):
                if template_tc is not None:
                    new_tc = copy.deepcopy(template_tc)
                    self._clear_row_green_tc(new_tc)
                    self._set_tc_text(new_tc, '')
                    row.append(new_tc)
        elif current > target_count:
            # Удаляем лишние ячейки
            for tc in tcs[target_count:]:
                row.remove(tc)

    def _clear_row_green(self, row):
        """Убирает зелёную заливку из всех ячеек строки."""
        for tc in row.findall(qn('w:tc')):
            self._clear_row_green_tc(tc)

    def _clear_row_green_tc(self, tc):
        """Убирает заливку из ячейки (устанавливает 'none')."""
        tcPr = tc.find(qn('w:tcPr'))
        if tcPr is None:
            return
        shd = tcPr.find(qn('w:shd'))
        if shd is not None:
            fill = shd.get(qn('w:fill'), '')
            # Убираем зелёную (B6D7A8) и синюю (C9DAF8) заливки
            if fill and fill.upper() not in ('', 'NONE', 'AUTO', 'FFFFFF', '000000'):
                shd.set(qn('w:fill'), 'none')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:val'), 'clear')

    def _set_tc_text(self, tc, text: str):
        """Устанавливает текст в ячейке таблицы."""
        paras = tc.findall(qn('w:p'))
        if paras:
            p = paras[0]
            # Удаляем все runs
            for r in p.findall(qn('w:r')):
                p.remove(r)
            # Добавляем новый run
            r = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = text
            if text and (text[0] == ' ' or text[-1] == ' '):
                t_elem.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            r.append(t_elem)
            p.append(r)
        else:
            p = OxmlElement('w:p')
            r = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = text
            r.append(t_elem)
            p.append(r)
            tc.append(p)

    def _set_tc_align(self, tc, align: str):
        """Устанавливает выравнивание текста в ячейке."""
        paras = tc.findall(qn('w:p'))
        for p in paras:
            pPr = p.find(qn('w:pPr'))
            if pPr is None:
                pPr = OxmlElement('w:pPr')
                p.insert(0, pPr)
            jc = pPr.find(qn('w:jc'))
            if jc is None:
                jc = OxmlElement('w:jc')
                pPr.append(jc)
            jc.set(qn('w:val'), align)

    def _update_signatures(self, doc, dir_title_raw: str, dir_name_raw: str,
                            customer_shortname: str = ''):
        """
        Обновляет подписи заказчика во всех таблицах подписей (раздел 13, приложения).
        Заменяет /ФИО/ и название организации ТОЛЬКО в правой колонке (заказчик).
        Левая колонка (подрядчик, Гидросервис) не затрагивается.
        """
        dir_name_short = _fio_to_short(dir_name_raw)
        # Фамилия заказчика для поиска
        customer_surname = dir_name_raw.split()[0] if dir_name_raw.split() else ''

        def _replace_in_cell(cell):
            """Replace /FIO/, org name, and director title in a single cell (customer side)."""
            for p in cell.paragraphs:
                text = p.text.strip()
                # Замена /ФИО/ подписи
                if '/' in text and text.count('/') >= 2:
                    inner = text.strip('/').strip()
                    # Подпись подрядчика (Гидросервис) — Сухов в любом месте строки
                    is_contractor = 'Сухов' in inner
                    if is_contractor:
                        continue
                    has_customer = bool(customer_surname and customer_surname in inner)
                    has_any_fio = bool(re.search(r'[A-ZА-Я][a-zа-я]+', inner))
                    if has_customer or has_any_fio:
                        # Сохраняем подчёркивания перед слешами
                        prefix = text[:text.index('/')]
                        _replace_paragraph_text(p, f'{prefix}/{dir_name_short}/')
                # Замена названия организации заказчика (ГБУЗ МО «Санаторий Пушкино» → новое название)
                elif customer_shortname and text and not text.startswith('«') and 'Гидросервис' not in text:
                    # Это строка с названием организации (не заголовок «ЗАКАЗЧИК», не подпись, не Гидросервис)
                    # Проверяем: если строка похожа на название организации (содержит «» или ГБУЗ/ООО/АО)
                    is_org_name = ('«' in text and '»' in text) or \
                                  any(prefix in text.upper() for prefix in ['ГБУЗ', 'ООО', 'АО', 'ПАО', 'ЗАО', 'МУП', 'ГУП', 'ФГУП'])
                    if is_org_name:
                        _replace_paragraph_text(p, customer_shortname)
                # Замена должности заказчика
                elif text and ('директор' in text.lower() or 'руководител' in text.lower()):
                    # Не трогаем если это подрядчик (Генеральный директор)
                    if 'Генеральный' not in text:
                        _replace_paragraph_text(p, f'{dir_title_raw.capitalize()}:')

        # Обходим все таблицы документа (пропускаем таблицу 0 — реквизиты, и таблицу 1 — смета)
        for ti, tbl in enumerate(doc.tables):
            if ti <= 1:
                continue  # Таблица 0 (реквизиты) и 1 (смета) заполняются отдельно
            for row in tbl.rows:
                cells = row.cells
                n = len(cells)
                if n >= 2:
                    # Двухколоночная таблица подписей:
                    # левая колонка (cells[0]) = ПОДРЯДЧИК (Гидросервис) — НЕ трогаем
                    # правая колонка (cells[-1]) = ЗАКАЗЧИК — заменяем
                    _replace_in_cell(cells[-1])
                elif n == 1:
                    # Одноколоночная строка — заменяем если есть фамилия заказчика
                    _replace_in_cell(cells[0])

    def _set_landscape_for_app3(self, doc):
        """
        Устанавливает альбомную ориентацию только для листов Приложения № 3.
        Остальные страницы документа остаются книжными.
        Реализация:
          1. Секция до Приложения 3 заканчивается книжным sectPr.
          2. Секция Приложения 3 имеет альбомный sectPr в параграфе перед следующей секцией.
          3. Секция после Приложения 3 возвращается к книжной ориентации.
        """
        body = doc.element.body
        children = list(body)

        # ---------------------------------------------------------------
        # Шаг 1: находим параграф с sectPr, предшествующий "Приложению № 3"
        # (это граница между Приложением 2 и Приложением 3)
        # ---------------------------------------------------------------
        app3_start_idx = None   # индекс первого элемента Приложения 3
        for i, child in enumerate(children):
            if child.tag == qn('w:p'):
                text = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
                if 'Приложение № 3' in text or 'Приложение №3' in text:
                    app3_start_idx = i
                    break

        if app3_start_idx is None:
            return

        # Ищем параграф с sectPr вне таблиц сразу перед Приложением 3
        # (идём назад от app3_start_idx)
        pre_app3_sect_p = None
        for i in range(app3_start_idx - 1, -1, -1):
            child = children[i]
            if child.tag == qn('w:p'):
                pPr = child.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                    pre_app3_sect_p = child
                    break

        # Если параграфа с sectPr нет перед Приложением 3 — создаём его
        if pre_app3_sect_p is None:
            pre_app3_sect_p = OxmlElement('w:p')
            new_pPr = OxmlElement('w:pPr')
            pre_app3_sect_p.append(new_pPr)
            children[app3_start_idx - 1].addnext(pre_app3_sect_p)
            children = list(body)  # обновляем список
            # Пересчитываем app3_start_idx
            for i, child in enumerate(children):
                if child.tag == qn('w:p'):
                    text = ''.join(t.text or '' for t in child.findall('.//' + qn('w:t')))
                    if 'Приложение № 3' in text or 'Приложение №3' in text:
                        app3_start_idx = i
                        break

        # Добавляем/обновляем sectPr в параграфе перед Приложением 3
        pPr_pre = pre_app3_sect_p.find(qn('w:pPr'))
        if pPr_pre is None:
            pPr_pre = OxmlElement('w:pPr')
            pre_app3_sect_p.insert(0, pPr_pre)
        sectPr_pre = pPr_pre.find(qn('w:sectPr'))
        if sectPr_pre is None:
            sectPr_pre = OxmlElement('w:sectPr')
            pPr_pre.append(sectPr_pre)

        # Устанавливаем книжную ориентацию для секции ДО Приложения 3
        pgSz_pre = sectPr_pre.find(qn('w:pgSz'))
        if pgSz_pre is None:
            pgSz_pre = OxmlElement('w:pgSz')
            sectPr_pre.append(pgSz_pre)
        pgSz_pre.set(qn('w:w'), '11906')
        pgSz_pre.set(qn('w:h'), '16838')
        pgSz_pre.attrib.pop(qn('w:orient'), None)  # portrait = по умолчанию

        # ---------------------------------------------------------------
        # Шаг 2: находим параграф с sectPr, замыкающий секцию Приложения 3
        # (это граница между Приложением 3 и следующим разделом)
        # Ищем параграф с sectPr после таблицы графика работ (таблица Приложения 3)
        # ---------------------------------------------------------------
        # Находим таблицу графика работ (содержит "Наименование" или "наименование")
        schedule_tbl_element = None
        for child in children:
            if child.tag == qn('w:tbl'):
                for row in child.findall('.//' + qn('w:tr'))[:2]:
                    for cell in row.findall('.//' + qn('w:tc')):
                        cell_text = ''.join(t.text or '' for t in cell.findall('.//' + qn('w:t')))
                        if 'Наименование' in cell_text or 'наименование' in cell_text:
                            schedule_tbl_element = child
                            break
                    if schedule_tbl_element:
                        break
            if schedule_tbl_element:
                break

        # Находим параграф с sectPr после таблицы графика
        post_app3_sect_p = None
        if schedule_tbl_element is not None:
            tbl_idx = list(body).index(schedule_tbl_element)
            for i in range(tbl_idx + 1, len(list(body))):
                child = list(body)[i]
                if child.tag == qn('w:p'):
                    pPr = child.find(qn('w:pPr'))
                    if pPr is not None and pPr.find(qn('w:sectPr')) is not None:
                        post_app3_sect_p = child
                        break
                    elif pPr is not None:
                        # Первый параграф после таблицы — используем его
                        post_app3_sect_p = child
                        break
                    else:
                        post_app3_sect_p = child
                        break

        # Если не нашли — создаём параграф с sectPr после таблицы
        if post_app3_sect_p is None:
            post_app3_sect_p = OxmlElement('w:p')
            new_pPr2 = OxmlElement('w:pPr')
            post_app3_sect_p.append(new_pPr2)
            if schedule_tbl_element is not None:
                schedule_tbl_element.addnext(post_app3_sect_p)
            else:
                body.append(post_app3_sect_p)

        # Добавляем/обновляем sectPr в параграфе после таблицы графика
        pPr_post = post_app3_sect_p.find(qn('w:pPr'))
        if pPr_post is None:
            pPr_post = OxmlElement('w:pPr')
            post_app3_sect_p.insert(0, pPr_post)
        sectPr_post = pPr_post.find(qn('w:sectPr'))
        if sectPr_post is None:
            sectPr_post = OxmlElement('w:sectPr')
            pPr_post.append(sectPr_post)

        # Устанавливаем альбомную ориентацию для секции Приложения 3
        pgSz_post = sectPr_post.find(qn('w:pgSz'))
        if pgSz_post is None:
            pgSz_post = OxmlElement('w:pgSz')
            sectPr_post.append(pgSz_post)
        pgSz_post.set(qn('w:w'), '16838')
        pgSz_post.set(qn('w:h'), '11906')
        pgSz_post.set(qn('w:orient'), 'landscape')

        # Уменьшаем поля для альбомной страницы, чтобы таблица влезла
        pgMar_post = sectPr_post.find(qn('w:pgMar'))
        if pgMar_post is None:
            pgMar_post = OxmlElement('w:pgMar')
            sectPr_post.append(pgMar_post)
        pgMar_post.set(qn('w:top'), '567')
        pgMar_post.set(qn('w:right'), '567')
        pgMar_post.set(qn('w:bottom'), '567')
        pgMar_post.set(qn('w:left'), '567')

        # ---------------------------------------------------------------
        # Шаг 3: находим sectPr в конце документа (финальный sectPr в body)
        # и убеждаемся, что он книжный
        # ---------------------------------------------------------------
        final_sectPr = body.find(qn('w:sectPr'))
        if final_sectPr is not None:
            pgSz_final = final_sectPr.find(qn('w:pgSz'))
            if pgSz_final is None:
                pgSz_final = OxmlElement('w:pgSz')
                final_sectPr.append(pgSz_final)
            pgSz_final.set(qn('w:w'), '11906')
            pgSz_final.set(qn('w:h'), '16838')
            pgSz_final.attrib.pop(qn('w:orient'), None)  # portrait
